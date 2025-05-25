"""Parsowanie CV (PDF / DOCX / TXT) do profilu kandydata."""
from __future__ import annotations

import io
import re

from . import config, skills as skills_mod


def extract_text(filename: str, data: bytes) -> str:
    """Wyciąga surowy tekst z pliku CV na podstawie rozszerzenia."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="ignore")
    # Próba jako tekst.
    return data.decode("utf-8", errors="ignore")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _from_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _detect_years(text: str) -> int:
    """Szacuje liczbę lat doświadczenia z fraz typu '5 lat doświadczenia'."""
    candidates = []
    for m in re.finditer(
        r"(\d{1,2})\+?\s*(?:lat|years?|roku|lata|yrs?)", text, re.IGNORECASE
    ):
        candidates.append(int(m.group(1)))
    # Odfiltruj absurdy.
    candidates = [c for c in candidates if 0 < c <= 40]
    return max(candidates) if candidates else 0


def _detect_seniority(text: str) -> int:
    lowered = text.lower()
    best = -1
    for kw, level in skills_mod.SENIORITY_KEYWORDS.items():
        if re.search(r"(?<!\w)" + re.escape(kw) + r"(?!\w)", lowered):
            best = max(best, level)
    if best >= 0:
        return best
    # Fallback po latach doświadczenia.
    years = _detect_years(text)
    if years >= 6:
        return 3
    if years >= 3:
        return 2
    if years >= 1:
        return 1
    return 1


def _detect_language(text: str) -> str:
    """Bardzo prosta detekcja PL/EN po typowych słowach."""
    pl_markers = ["doświadczenie", "umiejętności", "wykształcenie", "język", "obecnie"]
    score = sum(1 for w in pl_markers if w in text.lower())
    return "pl" if score >= 2 else "en"


def build_profile(filename: str, data: bytes) -> dict:
    """Buduje słownik profilu z pliku CV."""
    text = extract_text(filename, data)
    detected = sorted(skills_mod.extract_skills(text))
    return {
        "filename": filename,
        "skills": detected,
        "years": _detect_years(text),
        "seniority": _detect_seniority(text),
        "language": _detect_language(text),
        "text": text[:config.CV_TEXT_MAX_CHARS],  # przycinamy na potrzeby AI / podglądu
        "char_count": len(text),
    }


SENIORITY_LABELS = {0: "Intern", 1: "Junior", 2: "Mid", 3: "Senior", 4: "Lead/Architect"}
